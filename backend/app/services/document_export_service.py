from __future__ import annotations

import re
from io import BytesIO


class DocumentExportService:
    def render(self, *, title: str, markdown: str, format_name: str) -> tuple[bytes, str, str]:
        normalized_format = (format_name or "").strip().lower()
        if normalized_format == "md":
            return markdown.encode("utf-8"), "text/markdown; charset=utf-8", "md"
        if normalized_format == "docx":
            return self._render_docx(markdown=markdown), (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ), "docx"
        if normalized_format == "pdf":
            return self._render_pdf(title=title, markdown=markdown), "application/pdf", "pdf"
        raise ValueError("Unsupported export format")

    def _render_docx(self, *, markdown: str) -> bytes:
        try:
            from docx import Document  # type: ignore
        except Exception as e:
            raise RuntimeError("DOCX export is unavailable. Install python-docx.") from e

        doc = Document()
        lines = markdown.splitlines()
        i = 0
        in_code_block = False
        code_buffer: list[str] = []

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if stripped.startswith("```"):
                if in_code_block:
                    if code_buffer:
                        p = doc.add_paragraph("\n".join(code_buffer))
                        for run in p.runs:
                            run.font.name = "Courier New"
                    code_buffer = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_buffer.append(line)
                i += 1
                continue

            if not stripped:
                i += 1
                continue

            heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
            if heading_match:
                level = min(6, len(heading_match.group(1)))
                text = heading_match.group(2).strip()
                doc.add_heading(text, level=level)
                i += 1
                continue

            if self._is_table_header(lines, i):
                table_lines = [lines[i]]
                j = i + 2
                while j < len(lines) and "|" in lines[j]:
                    table_lines.append(lines[j])
                    j += 1
                self._add_docx_table(doc, table_lines)
                i = j
                continue

            unordered_match = re.match(r"^\s*[-*+]\s+(.+)$", stripped)
            if unordered_match:
                doc.add_paragraph(self._clean_inline_markdown(unordered_match.group(1)), style="List Bullet")
                i += 1
                continue

            ordered_match = re.match(r"^\s*\d+\.\s+(.+)$", stripped)
            if ordered_match:
                doc.add_paragraph(self._clean_inline_markdown(ordered_match.group(1)), style="List Number")
                i += 1
                continue

            quote_match = re.match(r"^\s*>\s?(.*)$", stripped)
            if quote_match:
                doc.add_paragraph(self._clean_inline_markdown(quote_match.group(1)), style="Intense Quote")
                i += 1
                continue

            doc.add_paragraph(self._clean_inline_markdown(stripped))
            i += 1

        output = BytesIO()
        doc.save(output)
        return output.getvalue()

    def _render_pdf(self, *, title: str, markdown: str) -> bytes:
        try:
            from reportlab.lib.pagesizes import LETTER  # type: ignore
            from reportlab.pdfgen import canvas  # type: ignore
        except Exception as e:
            raise RuntimeError("PDF export is unavailable. Install reportlab.") from e

        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=LETTER)
        width, height = LETTER

        y = height - 56
        c.setFont("Helvetica-Bold", 16)
        c.drawString(56, y, title[:120])
        y -= 30

        c.setFont("Helvetica", 11)
        for line in self._markdown_to_text_lines(markdown):
            if y < 52:
                c.showPage()
                y = height - 56
                c.setFont("Helvetica", 11)
            c.drawString(56, y, line[:140])
            y -= 15

        c.save()
        return buffer.getvalue()

    def _is_table_header(self, lines: list[str], idx: int) -> bool:
        if idx + 1 >= len(lines):
            return False
        header = lines[idx]
        separator = lines[idx + 1]
        if "|" not in header:
            return False
        return bool(
            re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", separator.strip())
        )

    def _split_md_row(self, line: str) -> list[str]:
        row = line.strip()
        if row.startswith("|"):
            row = row[1:]
        if row.endswith("|"):
            row = row[:-1]
        return [self._clean_inline_markdown(cell.strip()) for cell in row.split("|")]

    def _add_docx_table(self, doc, table_lines: list[str]) -> None:
        rows = [self._split_md_row(line) for line in table_lines if "|" in line]
        if not rows:
            return
        column_count = max(len(r) for r in rows)
        table = doc.add_table(rows=len(rows), cols=column_count)
        table.style = "Table Grid"
        for r_idx, row in enumerate(rows):
            for c_idx in range(column_count):
                value = row[c_idx] if c_idx < len(row) else ""
                cell = table.cell(r_idx, c_idx)
                cell.text = value
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

    def _clean_inline_markdown(self, text: str) -> str:
        cleaned = text
        cleaned = re.sub(r"`([^`]+)`", r"\1", cleaned)
        cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", cleaned)
        cleaned = re.sub(r"\*([^*]+)\*", r"\1", cleaned)
        cleaned = re.sub(r"__([^_]+)__", r"\1", cleaned)
        cleaned = re.sub(r"_([^_]+)_", r"\1", cleaned)
        cleaned = re.sub(r"~~([^~]+)~~", r"\1", cleaned)
        cleaned = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", cleaned)
        return cleaned.strip()

    def _markdown_to_text_lines(self, markdown: str) -> list[str]:
        lines: list[str] = []
        in_code_block = False
        for raw_line in markdown.splitlines():
            line = raw_line.rstrip()
            stripped = line.strip()
            if stripped.startswith("```"):
                in_code_block = not in_code_block
                continue
            if in_code_block:
                lines.append(f"    {line}")
                continue
            if not stripped:
                lines.append("")
                continue
            heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
            if heading:
                lines.append(heading.group(2).upper())
                continue
            if re.match(r"^\s*[-*+]\s+(.+)$", stripped):
                content = re.sub(r"^\s*[-*+]\s+", "", stripped)
                lines.append(f"• {self._clean_inline_markdown(content)}")
                continue
            if re.match(r"^\s*\d+\.\s+(.+)$", stripped):
                content = re.sub(r"^\s*\d+\.\s+", "", stripped)
                lines.append(f"• {self._clean_inline_markdown(content)}")
                continue
            if stripped.startswith(">"):
                lines.append(self._clean_inline_markdown(stripped.lstrip(">").strip()))
                continue
            if "|" in stripped:
                parts = [p.strip() for p in stripped.strip("|").split("|")]
                if any(parts):
                    lines.append(" | ".join(self._clean_inline_markdown(p) for p in parts))
                continue
            lines.append(self._clean_inline_markdown(stripped))
        return lines
